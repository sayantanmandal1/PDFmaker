"""
Integration test for export service using the actual PostgreSQL database.
Run this after creating a project with generated content.
"""
import sys
from sqlalchemy.orm import Session
from database import SessionLocal, engine
from models.user import User
from models.project import Project
from services.export_service import ExportService
from services.project_service import ProjectService
from services.content_service import ContentService
from services.auth_service import AuthService
from docx import Document
from pptx import Presentation
from io import BytesIO


def test_export_with_real_db():
    """Test export functionality with the actual database."""
    db = SessionLocal()
    
    try:
        # Create a test user
        print("Creating test user...")
        test_email = f"export_test_{id(db)}@example.com"
        user = AuthService.register_user(
            db=db,
            email=test_email,
            password="TestPassword123!",
            name="Export Test User"
        )
        print(f"✓ Created user: {user.email}")
        
        # Create a Word project
        print("\nCreating Word project...")
        word_project = ProjectService.create_project(
            db=db,
            user_id=user.id,
            name="Export Test Word Document",
            document_type="word",
            topic="Testing Word Export Functionality"
        )
        print(f"✓ Created Word project: {word_project.id}")
        
        # Add sections with content
        print("Adding sections with content...")
        ContentService.create_section(
            db=db,
            project_id=word_project.id,
            header="Introduction",
            content="This is the introduction section.\nIt contains multiple paragraphs.\nEach paragraph is on a new line.",
            position=0
        )
        ContentService.create_section(
            db=db,
            project_id=word_project.id,
            header="Main Content",
            content="This is the main content section with detailed information.",
            position=1
        )
        ContentService.create_section(
            db=db,
            project_id=word_project.id,
            header="Conclusion",
            content="This is the conclusion section that wraps up the document.",
            position=2
        )
        print("✓ Added 3 sections with content")
        
        # Test Word export
        print("\nTesting Word document export...")
        file_stream, filename = ExportService.export_word_document(db, word_project.id)
        print(f"✓ Export successful: {filename}")
        print(f"  File size: {len(file_stream.getvalue())} bytes")
        
        # Verify the document can be opened
        file_stream.seek(0)
        doc = Document(file_stream)
        print(f"  Document has {len(doc.paragraphs)} paragraphs")
        print(f"  Document has {len([p for p in doc.paragraphs if p.style.name.startswith('Heading')])} headings")
        
        # Create a PowerPoint project
        print("\nCreating PowerPoint project...")
        ppt_project = ProjectService.create_project(
            db=db,
            user_id=user.id,
            name="Export Test PowerPoint",
            document_type="powerpoint",
            topic="Testing PowerPoint Export Functionality"
        )
        print(f"✓ Created PowerPoint project: {ppt_project.id}")
        
        # Add slides with content
        print("Adding slides with content...")
        ContentService.create_slide(
            db=db,
            project_id=ppt_project.id,
            title="Introduction",
            content="Welcome to the presentation\nKey objectives\nOverview of topics",
            position=0
        )
        ContentService.create_slide(
            db=db,
            project_id=ppt_project.id,
            title="Main Topic",
            content="First point\nSecond point\nThird point",
            position=1
        )
        ContentService.create_slide(
            db=db,
            project_id=ppt_project.id,
            title="Conclusion",
            content="Summary\nNext steps\nThank you",
            position=2
        )
        print("✓ Added 3 slides with content")
        
        # Test PowerPoint export
        print("\nTesting PowerPoint presentation export...")
        file_stream, filename = ExportService.export_powerpoint_presentation(db, ppt_project.id)
        print(f"✓ Export successful: {filename}")
        print(f"  File size: {len(file_stream.getvalue())} bytes")
        
        # Verify the presentation can be opened
        file_stream.seek(0)
        prs = Presentation(file_stream)
        print(f"  Presentation has {len(prs.slides)} slides")
        
        # Clean up
        print("\nCleaning up test data...")
        ProjectService.delete_project(db, word_project.id, user.id)
        ProjectService.delete_project(db, ppt_project.id, user.id)
        db.delete(user)
        db.commit()
        print("✓ Cleanup complete")
        
        print("\n" + "=" * 60)
        print("✓ ALL EXPORT TESTS PASSED")
        print("=" * 60)
        
    except Exception as e:
        print(f"\n✗ Test failed: {str(e)}")
        import traceback
        traceback.print_exc()
        db.rollback()
        sys.exit(1)
    finally:
        db.close()


if __name__ == "__main__":
    print("=" * 60)
    print("Export Service Integration Test")
    print("=" * 60)
    print("\nThis test uses the actual PostgreSQL database")
    print("It will create test data and clean it up afterwards\n")
    
    test_export_with_real_db()
